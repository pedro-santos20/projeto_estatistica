import pandas as pd
import sys
import os

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx import Document

from python.analise_descritiva.distribuicao_unidades_estado import analisar_distribuicao_unidade
from python.analise_descritiva.correlacao_latitute_longitude_unidades import correlacao_lat_long
from python.analise_descritiva.distribuicao_tipo_gestao import analisar_distribuicao_tipo_gestao
from python.analise_descritiva.distribuicao_unidades_juridica import analisar_distribuicao_juridica
from python.analise_descritiva.distribuicao_unidades_tipo_atendimento import analisar_distribuicao_tipo_atendimento
from python.analise_descritiva.media_estabelecimentos_servicos_especificos import media_servicos_especificos
from python.inferencia_estatistica.diferenca_publica_privada_leitos import diferenca_unidade_publica_privada_leitos
from python.inferencia_estatistica.relacao_unidades_populacao import relacao_unidades_populacao
from python.probabilidade.probabilidade_unidade_area_urbana import probabilidade_area_urbana

data = pd.read_csv('/home/pedro/projeto_estatistica/cnes_estabelecimentos.csv', sep=';', encoding='ISO-8859-1', low_memory=False)

doc = Document()
doc.add_heading('Relatório de Análise Exploratória de Dados', 0)

def adicionar_analise(titulo, func):
    doc.add_heading(titulo, level=1)
    resultado = func(data)
    doc.add_paragraph(str(resultado))

adicionar_analise("Distribuição das Unidades de Saúde por Estado", analisar_distribuicao_unidade)
adicionar_analise("Correlação entre Latitude e Longitude das Unidades", correlacao_lat_long)
adicionar_analise("Distribuição das Unidades por Tipo de Gestão", analisar_distribuicao_tipo_gestao)
adicionar_analise("Distribuição das Unidades por Natureza Jurídica", analisar_distribuicao_juridica)
adicionar_analise("Distribuição das Unidades por Tipo de Atendimento", analisar_distribuicao_tipo_atendimento)
adicionar_analise("Média de Estabelecimentos com Serviços Específicos", media_servicos_especificos)
adicionar_analise("Diferença no Número de Leitos entre Unidades Públicas e Privadas", diferenca_unidade_publica_privada_leitos)
adicionar_analise("Relação entre Número de Unidades e População", relacao_unidades_populacao)
adicionar_analise("Probabilidade de uma Unidade Estar em Área Urbana", probabilidade_area_urbana)

doc.save("relatorio_analise_dados.docx")

print("Relatório gerado com sucesso!")